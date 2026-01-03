"""
Medical Manuscript Formatter
============================

This script converts Markdown (.md) or Text (.txt) manuscript files into formatted 
Microsoft Word (.docx) documents, adhering to standard medical journal submission guidelines.

Key Features:
- **Page Layout**: 1-inch margins, Times New Roman 12pt font.
- **Line Numbering**: Continuous line numbering for manuscript files.
- **Page Numbering**: "PAGE X" format in the footer.
- **Document Types**:
    - **Main Manuscript**: Line numbers, double spacing, first line indent.
    - **Cover Letter**: No line numbers, single spacing, block paragraphs (no indent).
    - **Title Page**: No line numbers, double spacing, no indent.
    - **Highlights**: No line numbers, double spacing, no indent.
    - **Abstract (Standalone)**: No line numbers, double spacing, no indent.
- **Table Handling**: 
    - Converts Markdown tables to Word tables.
    - Automatically switches to Landscape orientation for wide tables (>6 columns).
    - Centers text in cells (horizontal and vertical).
    - Sets table width to 100% of page width with adaptive column widths (autofit).
    - Handles table titles and footnotes, keeping them on the same page as the table.
    - Correctly parses footnotes even if followed by horizontal rules.
- **Figure Legends**:
    - Bold "Figure X" titles and Panel identifiers ("A.", "B.").
    - Removes other bold formatting within the legend text to strictly follow guidelines.
    - No indentation.
    - Merges panel descriptions (A., B., etc.) into the main figure paragraph.
    - Adds spacing after each figure legend.
- **Abstract**: Allows bold text for structured abstract headings (Background, Methods, etc.).
- **Section Breaks**: Smartly manages section breaks to avoid blank pages between Abstract/Introduction.

Usage:
    python format_manuscript.py [input_path]

    If input_path is not provided, a file dialog will open.
"""

import re
import os
import argparse
import tkinter as tk
from tkinter import filedialog
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_ALIGN_VERTICAL

# =============================================================================
# Helper Functions for Document Formatting
# =============================================================================

def set_margins(doc):
    """Set 1 inch margins for all sections."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_line_numbering(doc):
    """Add continuous line numbering to the document."""
    sect = doc.sections[0]
    sect_pr = sect._sectPr
    line_num_type = OxmlElement('w:lnNumType')
    line_num_type.set(qn('w:countBy'), '1')
    line_num_type.set(qn('w:restart'), 'continuous')
    sect_pr.append(line_num_type)

def add_page_numbers(doc):
    """Add page numbers to the footer in 'PAGE X' format."""
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    
    # Add 'PAGE' field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    # Format footer font
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def set_font_style(run, bold=False, italic=False):
    """Set standard font style: Times New Roman, 12pt, Black."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0, 0, 0)
    run.bold = bold
    run.italic = italic

def format_paragraph(paragraph, style_type='Body', alignment=WD_ALIGN_PARAGRAPH.LEFT, doc_config=None):
    """Apply paragraph formatting based on the content type and document configuration."""
    if doc_config is None:
        doc_config = {
            'line_spacing_rule': WD_LINE_SPACING.DOUBLE,
            'first_line_indent': Inches(0.5),
            'space_after': Pt(0)
        }

    paragraph_format = paragraph.paragraph_format
    paragraph_format.line_spacing_rule = doc_config.get('line_spacing_rule', WD_LINE_SPACING.DOUBLE)
    paragraph_format.alignment = alignment
    
    if style_type == 'Body':
        paragraph_format.first_line_indent = doc_config.get('first_line_indent', Inches(0.5))
        paragraph_format.space_after = doc_config.get('space_after', Pt(0))
    elif style_type == 'Heading':
        paragraph_format.first_line_indent = Inches(0)
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)
    elif style_type == 'Title':
        paragraph_format.first_line_indent = Inches(0)
        paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif style_type == 'List':
        paragraph_format.left_indent = Inches(0.5)
        paragraph_format.first_line_indent = Inches(-0.25)
    elif style_type == 'Table':
        paragraph_format.first_line_indent = Inches(0)
        paragraph_format.left_indent = Inches(0)
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def set_cell_border(cell, **kwargs):
    """
    Set cell's border.
    Usage: set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"})
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def ensure_portrait(doc):
    """
    Ensure the current section is Portrait. 
    If it is Landscape, create a new Portrait section.
    Returns True if a new section was added, False otherwise.
    """
    if doc.sections and doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE:
        section = doc.add_section(WD_SECTION.NEW_PAGE)
        section.orientation = WD_ORIENT.PORTRAIT
        if section.page_width > section.page_height:
            section.page_width, section.page_height = section.page_height, section.page_width
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        return True
    return False

# =============================================================================
# Table Handling Functions
# =============================================================================

def create_three_line_table(doc, header, rows, title=None, footnotes=None):
    """
    Create a three-line table with automatic landscape orientation for wide tables.
    Ensures each table starts on a new page.
    Includes Title (before) and Footnotes (after) in the same section.
    """
    num_cols = len(header)
    use_landscape = num_cols > 6
    
    # Handle Page Orientation
    if use_landscape:
        if doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE:
            doc.add_page_break()
        else:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = WD_ORIENT.LANDSCAPE
            if section.page_width < section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
    else:
        if doc.sections[-1].orientation == WD_ORIENT.LANDSCAPE:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            section.orientation = WD_ORIENT.PORTRAIT
            if section.page_width > section.page_height:
                section.page_width, section.page_height = section.page_height, section.page_width
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        else:
            doc.add_page_break()

    # Add Title
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        set_font_style(run, bold=True)
        p.paragraph_format.first_line_indent = Inches(0)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Create Table
    table = doc.add_table(rows=1, cols=len(header))
    table.autofit = True # Enable autofit for adaptive column widths
    
    # Set table width to 100%
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')
    
    # Header Row
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(text)
        set_font_style(run, bold=True)
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        format_paragraph(p, style_type='Table', alignment=WD_ALIGN_PARAGRAPH.CENTER)
        
        # Top Border (Thick)
        set_cell_border(hdr_cells[i], top={"sz": 12, "val": "single", "color": "000000"})
        # Bottom Border (Thin)
        set_cell_border(hdr_cells[i], bottom={"sz": 6, "val": "single", "color": "000000"})

    # Data Rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(row_data):
            if i < len(row_cells):
                p = row_cells[i].paragraphs[0]
                run = p.add_run(text)
                set_font_style(run)
                row_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                format_paragraph(p, style_type='Table', alignment=WD_ALIGN_PARAGRAPH.CENTER)
    
    # Bottom Border of Last Row (Thick)
    last_row_cells = table.rows[-1].cells
    for cell in last_row_cells:
        set_cell_border(cell, bottom={"sz": 12, "val": "single", "color": "000000"})

    # Add Footnotes
    if footnotes:
        for note in footnotes:
            p = doc.add_paragraph()
            run = p.add_run(note)
            set_font_style(run)
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

def parse_markdown_table(lines, start_index):
    """Parse a Markdown table from lines."""
    # Strip whitespace then strip pipes to handle trailing pipes correctly
    header = [x.strip() for x in lines[start_index].strip().strip('|').split('|')]
    # Skip separator line (start_index + 1)
    rows = []
    i = start_index + 2
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        row = [x.strip() for x in line.strip().strip('|').split('|')]
        rows.append(row)
        i += 1
    return header, rows, i - 1

# =============================================================================
# Main Processing Logic
# =============================================================================

def process_markdown_file(md_path, docx_path, doc_config):
    """
    Main function to process a Markdown file and generate a DOCX.
    """
    doc = Document()
    set_margins(doc)
    
    # Set default style
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.font.color.rgb = RGBColor(0, 0, 0)
    
    if doc_config['line_numbers']:
        add_line_numbering(doc)
    
    add_page_numbers(doc)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    i = 0
    pending_table_title = None
    in_figure_legends_section = False
    in_abstract_section = False
    current_figure_paragraph = None
    
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:
            i += 1
            continue
            
        # Ignore horizontal rules
        if line == '---' or line == '***':
            i += 1
            continue

        # ---------------------------------------------------------------------
        # Table Detection
        # ---------------------------------------------------------------------
        if line.startswith('|') and i + 1 < len(lines) and '---' in lines[i+1]:
            header, rows, end_index = parse_markdown_table(lines, i)
            
            # Look ahead for footnotes
            footnotes = []
            next_idx = end_index + 1
            while next_idx < len(lines):
                next_line = lines[next_idx].strip()
                if not next_line:
                    next_idx += 1
                    continue
                # Stop if we hit a header or another table or list or horizontal rule
                if next_line.startswith('#') or next_line.startswith('|') or next_line.startswith('* ') or next_line.startswith('- ') or next_line == '---' or next_line == '***':
                    break
                # Assume it's a footnote
                footnotes.append(next_line)
                next_idx += 1
            
            create_three_line_table(doc, header, rows, title=pending_table_title, footnotes=footnotes)
            pending_table_title = None # Consumed
            i = next_idx
            continue
            
        # ---------------------------------------------------------------------
        # Headings
        # ---------------------------------------------------------------------
        if line.startswith('#'):
            level = len(line.split()[0])
            text = line.lstrip('#').strip()
            
            # Check if this heading is actually a Table Title
            if re.match(r'^Table\s+\d+', text):
                pending_table_title = text
                i += 1
                continue

            # Ensure we are back in portrait mode for text
            switched_to_portrait = ensure_portrait(doc)
            
            # Check if we are entering Figure Legends or Abstract
            if text.lower() == 'figure legends':
                in_figure_legends_section = True
                in_abstract_section = False
            elif text.lower() == 'abstract':
                in_abstract_section = True
                in_figure_legends_section = False
            else:
                in_figure_legends_section = False
                in_abstract_section = False
                current_figure_paragraph = None
            
            # Add Page Breaks before specific sections
            # Only for Manuscript file
            if doc_config['line_numbers']: # Use line_numbers as proxy for "is_manuscript"
                # Check for major sections
                # Note: 'Abstract' might be early, but 'Introduction', 'References', 'Figure Legends' definitely need breaks
                # Removed 'introduction' from forced page break list to avoid potential blank page issues if user requests
                if text.lower() in ['abstract', 'references', 'figure legends', 'tables', 'figures']:
                    # Only add page break if we didn't just switch orientation (which adds a page break)
                    if not switched_to_portrait:
                        doc.add_page_break()
            
            p = doc.add_paragraph()
            run = p.add_run(text)
            set_font_style(run, bold=True) # Headings bold
            format_paragraph(p, style_type='Heading', doc_config=doc_config)
            i += 1
            continue
            
        # ---------------------------------------------------------------------
        # List Items
        # ---------------------------------------------------------------------
        if line.startswith('* ') or line.startswith('- '):
            ensure_portrait(doc)
            text = line[2:].strip()
            p = doc.add_paragraph()
            
            # Add bullet manually
            run = p.add_run('•\t')
            set_font_style(run)
            
            # Handle bold/italic in text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    set_font_style(run, bold=True)
                else:
                    run = p.add_run(part)
                    set_font_style(run)
            format_paragraph(p, style_type='List', doc_config=doc_config)
            i += 1
            continue
            
        # ---------------------------------------------------------------------
        # Normal Paragraph
        # ---------------------------------------------------------------------
        # Check if this is a Table Title (starts with "Table \d")
        if re.match(r'^Table\s+\d+', line):
            pending_table_title = line
            i += 1
            continue
            
        # If we had a pending table title but the next line wasn't a table, flush it
        if pending_table_title:
            ensure_portrait(doc)
            p = doc.add_paragraph()
            run = p.add_run(pending_table_title)
            set_font_style(run, bold=True)
            format_paragraph(p, style_type='Body', doc_config=doc_config)
            pending_table_title = None
            # Don't increment i, process current line again
            continue

        ensure_portrait(doc)
        
        # Figure Legend Special Handling
        if in_figure_legends_section:
            # Check if it is a Figure Title (starts with Figure X or **Figure X)
            # The markdown has **Figure 1: ...** so line starts with **Figure
            is_figure_title = re.match(r'^(\*\*)?Figure\s+\d+', line)
            
            if is_figure_title:
                # Start new paragraph for Figure Title
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Inches(0)
                p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
                p.paragraph_format.space_after = Pt(12) # Add spacing after figure legend
                current_figure_paragraph = p
                
                # Process tokens
                tokens = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|\^.*?\^)', line)
                for token in tokens:
                    if not token: continue
                    if token.startswith('**') and token.endswith('**'):
                        content = token[2:-2]
                        run = p.add_run(content)
                        # Only bold if it matches "Figure X" pattern
                        if re.match(r'^Figure\s+\d+', content):
                            set_font_style(run, bold=True)
                        else:
                            set_font_style(run, bold=False)
                    elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
                        # Italic case
                        content = token[1:-1]
                        run = p.add_run(content)
                        set_font_style(run, italic=True)
                    elif token.startswith('^') and token.endswith('^'):
                        run = p.add_run(token[1:-1])
                        set_font_style(run)
                        run.font.superscript = True
                    else:
                        run = p.add_run(token)
                        set_font_style(run)
                i += 1
                continue
            
            elif current_figure_paragraph and re.match(r'^(\*\*)?[A-Z]\.(\*\*)?', line):
                # This is a panel description (A., B., etc)
                # Append to current paragraph
                p = current_figure_paragraph
                p.add_run(' ') # Add space before panel
                
                # Process tokens
                tokens = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|\^.*?\^)', line)
                for token in tokens:
                    if not token: continue
                    if token.startswith('**') and token.endswith('**'):
                        content = token[2:-2]
                        run = p.add_run(content)
                        # Only bold if it matches "A." pattern
                        if re.match(r'^[A-Z]\.$', content):
                            set_font_style(run, bold=True)
                        else:
                            set_font_style(run, bold=False)
                    elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
                        # Italic case
                        content = token[1:-1]
                        run = p.add_run(content)
                        set_font_style(run, italic=True)
                    elif token.startswith('^') and token.endswith('^'):
                        run = p.add_run(token[1:-1])
                        set_font_style(run)
                        run.font.superscript = True
                    else:
                        run = p.add_run(token)
                        set_font_style(run)
                i += 1
                continue

        p = doc.add_paragraph()
        
        # Check for Figure Legend Title (fallback if not in section, though we track section now)
        is_figure_legend_line = False
        if re.match(r'^Figure\s+\d+', line):
            is_figure_legend_line = True
        
        # Handle bold/italic/superscript in text
        # Simple parser for **bold**, *italic*, and ^super^
        # We need to split by all patterns
        # Regex explanation:
        # (\*\*.*?\*\*) : Matches **bold**
        # (\*[^*]+?\*)  : Matches *italic* (ensure no internal *)
        # (\^.*?\^)     : Matches ^super^
        tokens = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*|\^.*?\^)', line)
        
        for token in tokens:
            if not token: continue
            
            if token.startswith('**') and token.endswith('**'):
                content = token[2:-2]
                run = p.add_run(content)
                
                should_bold = doc_config.get('allow_bold', False)
                if in_abstract_section:
                    should_bold = True
                
                set_font_style(run, bold=should_bold)
            
            elif token.startswith('*') and token.endswith('*') and not token.startswith('**'):
                # Italic case
                content = token[1:-1]
                run = p.add_run(content)
                set_font_style(run, italic=True)
                
            elif token.startswith('^') and token.endswith('^'):
                run = p.add_run(token[1:-1])
                set_font_style(run)
                run.font.superscript = True
            else:
                run = p.add_run(token)
                set_font_style(run)
        
        if is_figure_legend_line:
            # No indent for Figure Legends
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.paragraph_format.space_after = Pt(12) # Add spacing after figure legend
        else:
            format_paragraph(p, style_type='Body', doc_config=doc_config)
            
        i += 1
        
    try:
        doc.save(docx_path)
        print(f"Generated: {docx_path}")
    except PermissionError:
        print(f"Error: Could not save to {docx_path}. The file might be open.")
        new_docx_path = docx_path.replace(".docx", "_v2.docx")
        print(f"Trying to save to {new_docx_path} instead...")
        doc.save(new_docx_path)
        print(f"Generated: {new_docx_path}")

def main():
    parser = argparse.ArgumentParser(description="Convert Markdown/Text manuscript to formatted Word DOCX.")
    parser.add_argument("input_path", nargs='?', default=None, help="Input file or directory path")
    args = parser.parse_args()

    input_path = args.input_path
    files_to_process = []

    # If no argument provided, open file dialog
    if input_path is None:
        print("No input path provided. Opening file dialog...")
        root = tk.Tk()
        root.withdraw() # Hide the main window
        file_paths = filedialog.askopenfilenames(
            title="Select Manuscript Files",
            filetypes=[("Markdown & Text", "*.md *.txt"), ("Markdown", "*.md"), ("Text", "*.txt"), ("All Files", "*.*")]
        )
        files_to_process = list(file_paths)
        if not files_to_process:
            print("No files selected. Exiting.")
            return
    else:
        # Handle command line argument
        if os.path.isfile(input_path):
            files_to_process = [input_path]
        elif os.path.isdir(input_path):
            print(f"Scanning directory: {input_path}")
            for filename in os.listdir(input_path):
                if filename.lower().endswith((".md", ".txt")):
                    files_to_process.append(os.path.join(input_path, filename))
        else:
            print("Error: Invalid path.")
            return

    # Process files
    for file_path in files_to_process:
        if not file_path.lower().endswith((".md", ".txt")):
            print(f"Skipping {file_path} (Not a .md or .txt file)")
            continue

        # Determine output path
        base, ext = os.path.splitext(file_path)
        docx_path = base + ".docx"
        
        # Heuristic to determine document type and configuration
        filename = os.path.basename(file_path).lower()
        
        # Default Configuration (Main Manuscript)
        doc_config = {
            'line_numbers': True,
            'line_spacing_rule': WD_LINE_SPACING.DOUBLE,
            'first_line_indent': Inches(0.5),
            'space_after': Pt(0),
            'allow_bold': False # Strict no bold for body
        }
        
        if "cover_letter" in filename or "cover letter" in filename:
            print("Detected Cover Letter format.")
            doc_config = {
                'line_numbers': False,
                'line_spacing_rule': WD_LINE_SPACING.SINGLE, # Or 1.15/1.5 depending on preference, Single is standard for formal letters
                'first_line_indent': Inches(0),
                'space_after': Pt(12), # Block paragraph style
                'allow_bold': True
            }
        elif "title_page" in filename or "title page" in filename:
            print("Detected Title Page format.")
            doc_config = {
                'line_numbers': False,
                'line_spacing_rule': WD_LINE_SPACING.DOUBLE,
                'first_line_indent': Inches(0),
                'space_after': Pt(12), # Add spacing between elements for readability
                'allow_bold': True
            }
        elif "highlights" in filename:
            print("Detected Highlights format.")
            doc_config = {
                'line_numbers': False,
                'line_spacing_rule': WD_LINE_SPACING.DOUBLE,
                'first_line_indent': Inches(0),
                'space_after': Pt(0),
                'allow_bold': True
            }
        elif "abstract" in filename and not ("main" in filename or "manuscript" in filename):
             # Standalone Abstract
            print("Detected Standalone Abstract format.")
            doc_config = {
                'line_numbers': False,
                'line_spacing_rule': WD_LINE_SPACING.DOUBLE,
                'first_line_indent': Inches(0), # Often abstracts are block or indented, let's go with no indent for standalone
                'space_after': Pt(0),
                'allow_bold': True
            }
        
        print(f"Processing {filename}...")
        try:
            process_markdown_file(file_path, docx_path, doc_config)
        except Exception as e:
            print(f"Error processing {filename}: {e}")

if __name__ == "__main__":
    main()
